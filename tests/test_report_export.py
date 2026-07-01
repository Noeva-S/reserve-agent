from __future__ import annotations

from io import BytesIO
from pathlib import Path
import unittest

from docx import Document

from reserve_agent.agent.explanation import generate_data_diagnosis
from reserve_agent.data.loader import DataQualityReport
from reserve_agent.data.loader import load_excel_to_triangle
from reserve_agent.data_processing import load_exposure_data
from reserve_agent.explanation import generate_agent_explanation, generate_method_notes
from reserve_agent.exports import build_word_report
from reserve_agent.reserving import run_reserving_models


class WordReportTests(unittest.TestCase):
    def test_data_diagnosis_handles_empty_year_ranges(self) -> None:
        report = DataQualityReport(
            row_count=10,
            claim_count=0,
            accident_years=[],
            valuation_years=[],
            missing_values=0,
            negative_amount_cells=0,
            zero_claim_rows=0,
            notes=["当前工作表可能不是标准赔案建模数据。"],
        )

        messages = generate_data_diagnosis(report)

        self.assertTrue(any("未知" in message for message in messages))

    def test_complete_word_report_can_be_opened(self) -> None:
        project_root = Path(__file__).resolve().parents[1]
        workbook = project_root / "Chapter 13a - IBNR (triangle-based) v4.xlsx"
        result = load_excel_to_triangle(workbook, "2. Claims data", measure="Paid")
        outputs = run_reserving_models(result.triangle, load_exposure_data(workbook), 0.72)
        explanation = generate_agent_explanation(result.quality, outputs)

        report_bytes = build_word_report(
            source_file=workbook.name,
            sheet_name="2. Claims data",
            format_name=result.format_name,
            measure="Paid",
            triangle=result.triangle,
            outputs=outputs,
            quality=result.quality,
            explanation_text=explanation,
            method_notes=generate_method_notes(),
        )

        self.assertGreater(len(report_bytes), 50_000)
        document = Document(BytesIO(report_bytes))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("准备金评估完整报告", text)
        self.assertIn("累计赔款三角形", text)
        self.assertIn("Agent 分析", text)
        self.assertGreaterEqual(len(document.tables), 4)
        self.assertGreaterEqual(len(document.inline_shapes), 2)


if __name__ == "__main__":
    unittest.main()
