from __future__ import annotations

import sys
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from reserve_agent.data_processing import (  # noqa: E402
    build_cumulative_triangle,
    find_default_workbook,
    load_claims_snapshot,
    load_exposure_data,
    quality_report,
)
from reserve_agent.explanation import (  # noqa: E402
    generate_agent_explanation,
    generate_data_diagnosis,
    generate_method_notes,
    generate_result_summary,
)
from reserve_agent.reserving import run_reserving_models  # noqa: E402


OUTPUT = ROOT / "output"
FIGURES = OUTPUT / "figures"


def set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(10.5)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        s = doc.styles[style_name]
        s.font.name = "黑体"


def add_title(doc: Document, title: str, subtitle: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(18)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        r2.font.size = Pt(12)


def add_table_from_df(doc: Document, df: pd.DataFrame, max_rows: int = 12) -> None:
    shown = df.head(max_rows).copy()
    table = doc.add_table(rows=1, cols=len(shown.columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, col in enumerate(shown.columns):
        hdr[i].text = str(col)
    for _, row in shown.iterrows():
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if isinstance(value, float):
                cells[i].text = f"{value:,.2f}"
            else:
                cells[i].text = str(value)


def save_figures(comparison: pd.DataFrame) -> tuple[Path, Path]:
    FIGURES.mkdir(parents=True, exist_ok=True)

    reserve_cols = ["Chain Ladder Reserve", "ELR Reserve", "BF Reserve", "Selected Reserve"]
    reserve_path = FIGURES / "reserve_comparison.png"
    _draw_grouped_bar_chart(
        comparison,
        reserve_cols,
        reserve_path,
        title="Reserve Comparison by Method",
        y_label="Reserve (million)",
    )

    ultimate_path = FIGURES / "ultimate_trend.png"
    _draw_line_chart(
        comparison,
        ["Latest Cumulative", "Selected Ultimate"],
        ultimate_path,
        title="Latest Cumulative vs Selected Ultimate",
        y_label="Amount (million)",
    )

    return reserve_path, ultimate_path


def _font() -> ImageFont.ImageFont:
    return ImageFont.load_default()


def _draw_axes(draw: ImageDraw.ImageDraw, plot: tuple[int, int, int, int], title: str, y_label: str) -> None:
    font = _font()
    left, top, right, bottom = plot
    draw.rectangle([0, 0, 1200, 640], fill="white")
    draw.line([left, bottom, right, bottom], fill="#333333", width=2)
    draw.line([left, top, left, bottom], fill="#333333", width=2)
    draw.text((left, 24), title, fill="#111111", font=font)
    draw.text((left, 52), y_label, fill="#555555", font=font)


def _draw_y_ticks(draw: ImageDraw.ImageDraw, plot: tuple[int, int, int, int], max_value: float) -> None:
    font = _font()
    left, top, right, bottom = plot
    for i in range(6):
        value = max_value * i / 5
        y = bottom - (bottom - top) * i / 5
        draw.line([left - 5, y, right, y], fill="#E5E7EB", width=1)
        draw.text((12, y - 7), f"{value:,.0f}", fill="#555555", font=font)


def _draw_grouped_bar_chart(
    comparison: pd.DataFrame,
    columns: list[str],
    path: Path,
    title: str,
    y_label: str,
) -> None:
    img = Image.new("RGB", (1200, 640), "white")
    draw = ImageDraw.Draw(img)
    font = _font()
    plot = (88, 92, 1130, 520)
    _draw_axes(draw, plot, title, y_label)

    values = comparison[columns].fillna(0) / 1_000_000
    max_value = max(float(values.max().max()), 1.0) * 1.12
    _draw_y_ticks(draw, plot, max_value)

    colors = ["#2563EB", "#059669", "#DC2626", "#7C3AED"]
    years = comparison["Accident Year"].astype(str).tolist()
    left, top, right, bottom = plot
    group_width = (right - left) / len(years)
    bar_width = group_width / (len(columns) + 2)

    for gi, year in enumerate(years):
        group_left = left + gi * group_width
        for ci, col in enumerate(columns):
            value = float(values.iloc[gi][col])
            height = (bottom - top) * value / max_value
            x0 = group_left + (ci + 0.7) * bar_width
            x1 = x0 + bar_width * 0.82
            y0 = bottom - height
            draw.rectangle([x0, y0, x1, bottom], fill=colors[ci])
        draw.text((group_left + group_width * 0.25, bottom + 12), year, fill="#333333", font=font)

    legend_x = 760
    for idx, col in enumerate(columns):
        y = 22 + idx * 18
        draw.rectangle([legend_x, y, legend_x + 12, y + 12], fill=colors[idx])
        draw.text((legend_x + 18, y - 1), col, fill="#333333", font=font)

    img.save(path)


def _draw_line_chart(
    comparison: pd.DataFrame,
    columns: list[str],
    path: Path,
    title: str,
    y_label: str,
) -> None:
    img = Image.new("RGB", (1200, 640), "white")
    draw = ImageDraw.Draw(img)
    font = _font()
    plot = (88, 92, 1130, 520)
    _draw_axes(draw, plot, title, y_label)

    values = comparison[columns].fillna(0) / 1_000_000
    max_value = max(float(values.max().max()), 1.0) * 1.12
    _draw_y_ticks(draw, plot, max_value)

    colors = ["#2563EB", "#DC2626"]
    years = comparison["Accident Year"].astype(str).tolist()
    left, top, right, bottom = plot
    step = (right - left) / max(len(years) - 1, 1)

    for ci, col in enumerate(columns):
        points = []
        for idx, value in enumerate(values[col].tolist()):
            x = left + idx * step
            y = bottom - (bottom - top) * float(value) / max_value
            points.append((x, y))
        if len(points) > 1:
            draw.line(points, fill=colors[ci], width=3)
        for x, y in points:
            draw.ellipse([x - 4, y - 4, x + 4, y + 4], fill=colors[ci])

    for idx, year in enumerate(years):
        x = left + idx * step
        draw.text((x - 12, bottom + 12), year, fill="#333333", font=font)

    legend_x = 810
    for idx, col in enumerate(columns):
        y = 24 + idx * 20
        draw.line([legend_x, y + 6, legend_x + 20, y + 6], fill=colors[idx], width=3)
        draw.text((legend_x + 28, y), col, fill="#333333", font=font)

    img.save(path)


def build_report() -> Path:
    OUTPUT.mkdir(exist_ok=True)
    workbook = find_default_workbook(ROOT)
    claims = load_claims_snapshot(workbook, "Claims data")
    triangle = build_cumulative_triangle(claims, "Paid")
    exposure = load_exposure_data(workbook)
    report = quality_report(claims, triangle)
    outputs = run_reserving_models(triangle, exposure, 0.72)
    reserve_fig, ultimate_fig = save_figures(outputs.comparison)

    doc = Document()
    set_default_font(doc)
    add_title(doc, "基于 SOA/CAS 规范的准备金评估智能 Agent 系统设计", "非寿险精算课程小组作业初版报告")

    doc.add_heading("一、作业背景与系统目标", level=1)
    doc.add_paragraph(
        "非寿险公司的赔款具有发生、报告、定损、支付和结案之间存在时间差的特点。"
        "在财务报告日，保险公司需要估计已经发生但尚未完全支付的赔款责任，通常称为未决赔款准备金或 outstanding claim reserves。"
        "准备金估计直接影响偿付能力、利润确认、产品定价和经营决策，因此是非寿险精算工作中的核心环节。"
        "本作业按照课程要求，设计并实现一个轻量化准备金评估智能 Agent 系统。系统以 Excel 赔案数据为输入，自动完成数据导入、清洗诊断、赔付进展三角生成、模型运算、结果可视化和解释文本输出。"
    )
    doc.add_paragraph(
        "初版系统暂不接入外部大模型 API，而采用规则型 Agent 思路：将准备金评估流程拆分为可自动执行的模块，并根据数据质量指标、模型差异和事故年贡献自动生成分析意见。"
        "这种设计可以降低 API 费用、网络和密钥管理风险，同时为后续接入 DeepSeek 等大模型保留清晰接口。"
    )

    doc.add_heading("二、规范依据与参考资料", level=1)
    doc.add_paragraph(
        "系统设计参考 CAS 教材 Estimating Unpaid Claims Using Basic Techniques 以及 SOA ASTAM Outstanding Claims Reserves 学习资料。"
        "CAS 教材强调准备金评估前的信息收集、发展三角构建、发展法、期望赔款法、Bornhuetter-Ferguson 法以及模型结果评价。"
        "SOA 资料则从未决赔款准备金的定义、run-off triangle 的时间维度、确定性链梯法和随机模型扩展等角度给出了较系统的理论框架。"
        "本初版系统主要实现基础确定性模型，并在结果解释中提示模型假设和适用条件。"
    )
    doc.add_paragraph(
        "从精算规范角度看，一个准备金评估系统至少应做到：数据来源可追溯，关键假设可解释，模型过程可复核，结果披露包含不确定性提示。"
        "因此，本系统没有将 AI 简化为聊天接口，而是把 AI Agent 理解为能够自主执行准备金评估流程、识别数据问题、比较模型结果并输出解释的自动化程序。"
    )

    doc.add_heading("三、系统架构设计", level=1)
    doc.add_paragraph(
        "系统采用 Python + Streamlit 实现，代码按职责拆分为数据处理模块、准备金模型模块、解释模块和交互界面模块。"
        "数据处理模块负责读取 Excel、识别评估年列、展开赔案快照数据、按事故年和发展期生成累计赔付三角；"
        "模型模块负责计算发展因子、年龄到最终因子、最终赔款和准备金；"
        "解释模块负责根据规则生成数据诊断、模型推荐和结果摘要；"
        "Streamlit 页面负责提供上传、参数选择、表格展示和图表展示。"
    )
    doc.add_paragraph(
        "该架构的优点是低耦合、易演示、易扩展。后续如果需要接入 DeepSeek API，只需要新增 LLM 客户端，并在解释模块中增加“若存在 API key 则调用大模型，否则使用规则解释”的分支，不需要重写数据处理或模型计算代码。"
    )

    doc.add_heading("四、数据来源与预处理", level=1)
    doc.add_paragraph(
        f"初版使用文件 {workbook.name} 中的 Claims data 工作表。该工作表以赔案为单位，每个赔案包含 Paid、O/S 和 Incurred 三类度量，并按评估年份记录累计金额。"
        "系统默认采用 Paid 口径构建累计赔付三角，事故年为 Loss Year，发展期为 valuation year 减 Loss Year。"
    )
    for line in generate_data_diagnosis(report):
        doc.add_paragraph(line, style=None)

    doc.add_heading("五、赔付进展三角构建", level=1)
    doc.add_paragraph(
        "设 C_{i,j} 表示事故年 i 在发展期 j 的累计已付赔款。系统首先将宽表中的评估年份列转换为长表，计算 development = valuation year - accident year，"
        "再按事故年和发展期汇总赔款金额，最后透视为三角矩阵。对于尚未观测的右下角单元格，系统保留为空值，不参与发展因子估计。"
    )
    tri_display = triangle.copy()
    tri_display.index.name = "Accident Year"
    tri_display = tri_display.reset_index()
    add_table_from_df(doc, tri_display, max_rows=8)

    doc.add_heading("六、模型原理", level=1)
    notes = generate_method_notes()
    doc.add_paragraph("1. Chain Ladder 链梯法。")
    doc.add_paragraph(
        notes["Chain Ladder"]
        + " 年龄到年龄发展因子的体积加权估计为 f_j = sum_i C_{i,j+1} / sum_i C_{i,j}。"
        " 对于最新发展期为 k 的事故年，其最终赔款估计为 C_{i,k} 乘以从 k 到最终期的累计发展因子。"
    )
    doc.add_paragraph("2. Expected Loss Ratio 期望赔付率法。")
    doc.add_paragraph(
        notes["Expected Loss Ratio"]
        + " 初版系统使用暴露量表和可调期望赔付率参数形成先验最终赔款，并保证先验最终赔款不低于当前已观测累计赔款。"
    )
    doc.add_paragraph("3. Bornhuetter-Ferguson 法。")
    doc.add_paragraph(
        notes["Bornhuetter-Ferguson"]
        + " 设年龄到最终因子为 LDF，则已报告比例约为 1/LDF，未报告比例为 1 - 1/LDF。"
        " BF 准备金 = 先验最终赔款 × 未报告比例，BF 最终赔款 = 最新累计赔款 + BF 准备金。"
    )

    doc.add_heading("七、模型结果与可视化", level=1)
    doc.add_paragraph("下表为各事故年的模型结果对比。Selected 口径采用 Chain Ladder 与 BF 的均值，便于初版展示。")
    add_table_from_df(doc, outputs.comparison, max_rows=10)
    doc.add_paragraph("图 1 展示不同模型下各事故年的准备金差异。")
    doc.add_picture(str(reserve_fig), width=Inches(6.2))
    doc.add_paragraph("图 2 展示最新累计赔款与展示口径最终赔款的对比。")
    doc.add_picture(str(ultimate_fig), width=Inches(6.2))
    for line in generate_result_summary(outputs):
        doc.add_paragraph(line)

    doc.add_heading("八、代码实现说明", level=1)
    doc.add_paragraph(
        "代码目录 reserve_agent 中包含四个主要文件。data_processing.py 负责 Excel 读取、评估年份识别、累计三角构建和数据质量报告；"
        "reserving.py 实现 Chain Ladder、ELR、BF 和模型结果比较；"
        "explanation.py 实现规则型 Agent 的诊断和解释文本；"
        "app.py 是 Streamlit 交互界面，提供数据上传、参数选择、表格展示和图表展示。"
    )
    doc.add_paragraph(
        "这种拆分使得模型计算可以脱离界面独立测试，也便于后续把解释模块替换为 DeepSeek API。"
        "例如后续可以新增 llm_client.py，读取环境变量 DEEPSEEK_API_KEY，将 comparison 表、发展因子和数据诊断传入大模型，返回自然语言解释。"
    )

    doc.add_heading("九、AI Agent 功能设计", level=1)
    doc.add_paragraph(
        "初版系统的智能性体现在自动化流程与规则推理，而非外部 API 调用。系统能够自动判断数据是否存在空值、负赔款、零赔案记录和最末发展期观测不足等问题；"
        "能够根据准备金占比、模型差异和事故年贡献给出模型选择建议；"
        "能够自动生成面向报告的解释文本，说明数据质量、模型结果、主要风险点和后续复核方向。"
    )
    doc.add_paragraph(
        "后续增强版可以接入 DeepSeek API，实现自由问答和更自然的文字生成。为了控制风险，建议保留规则型解释作为 fallback：当 API key 缺失、余额不足或网络不可用时，系统仍然可以完整运行。"
    )

    doc.add_heading("十、挑战与解决方案", level=1)
    doc.add_paragraph(
        "第一，源数据不是标准三角，而是按赔案和评估年展开的快照表。解决方法是先识别评估年列，再将宽表转换成长表，并按事故年和发展期汇总。"
        "第二，较新事故年发展不成熟，链梯法结果对早期发展因子较敏感。解决方法是同时展示 BF 和 ELR，并在解释中提示成熟度风险。"
        "第三，课程要求融入 AI 元素，但 API 可能带来费用和稳定性问题。解决方法是先实现规则型 Agent 初版，保证可运行和可提交，再预留大模型接口作为后续扩展。"
    )
    doc.add_paragraph(
        "在作业过程中，AI 辅助主要用于梳理系统架构、解释参考资料、设计模型流程、生成代码初稿和组织报告结构。"
        "人工复核仍然很重要，尤其是模型假设、发展因子选择、异常赔案处理和最终报告表述，需要结合课程知识进行判断。"
    )

    doc.add_heading("十一、小组分工", level=1)
    doc.add_paragraph(
        "初版分工可按模块划分：成员 A 负责数据清洗和三角构建；成员 B 负责 Chain Ladder、ELR 和 BF 模型；"
        "成员 C 负责 Streamlit 界面与可视化；成员 D 负责 Agent 解释模块、报告撰写和材料汇总。"
        "如小组人数调整，可将测试、截图和使用说明单独分配。"
    )

    doc.add_heading("十二、总结", level=1)
    doc.add_paragraph(
        "本作业初版完成了从赔案数据到准备金评估结果的完整自动化流程。"
        "系统虽然暂未接入外部大模型 API，但已经具备轻量化 Agent 的基本特征：自主读取数据、执行模型、诊断问题、输出结果并生成解释。"
        "后续可以在不改变核心模型代码的基础上接入 DeepSeek API，进一步增强自然语言交互和报告生成能力。"
    )

    path = OUTPUT / "非寿险准备金评估智能Agent系统设计报告_初版.docx"
    doc.save(path)
    return path


def build_manual() -> Path:
    doc = Document()
    set_default_font(doc)
    add_title(doc, "非寿险准备金评估智能 Agent 系统使用说明书")
    doc.add_heading("一、运行环境", level=1)
    doc.add_paragraph("推荐使用 Python 3.10 及以上版本。需要安装 pandas、numpy、openpyxl、streamlit、plotly、python-docx、matplotlib。")
    doc.add_heading("二、启动步骤", level=1)
    doc.add_paragraph("1. 打开 PowerShell，进入大作业根目录。")
    doc.add_paragraph("2. 执行 python -m pip install -r reserve_agent/requirements.txt 安装依赖。")
    doc.add_paragraph("3. 执行 streamlit run reserve_agent/app.py 启动系统。")
    doc.add_paragraph("4. 浏览器打开命令行显示的本地地址，一般为 http://localhost:8501。")
    doc.add_heading("三、功能说明", level=1)
    doc.add_paragraph("左侧栏可以选择 Excel 数据文件、赔案数据工作表、三角口径和期望赔付率参数。")
    doc.add_paragraph("数据诊断页展示数据规模、空值、负金额、零赔案等提示。")
    doc.add_paragraph("赔付三角页展示累计赔付进展三角和发展因子。")
    doc.add_paragraph("模型结果页展示 Chain Ladder、ELR、BF 及展示口径的准备金对比。")
    doc.add_paragraph("可视化页展示模型准备金柱状图和最终赔款趋势图。")
    doc.add_paragraph("Agent 解释页输出可复制的规则型自动解释文本。")
    doc.add_heading("四、注意事项", level=1)
    doc.add_paragraph("初版默认读取 Chapter 08 - Data sets - Examples.xlsx 中的 Claims data 工作表。")
    doc.add_paragraph("如上传新文件，应尽量保持 Claim ID、Loss Year、Type 和评估年份列的结构一致。")
    doc.add_paragraph("系统支持可选 DeepSeek API 增强解释。若不启用 API，系统会自动使用规则型解释，不会产生调用费用。")
    doc.add_paragraph("如需启用 DeepSeek，可在项目根目录创建 .env 文件，写入 DEEPSEEK_API_KEY=你的Key，也可以在系统侧边栏密码框临时输入。不要把真实 API key 写入报告或提交材料。")
    path = OUTPUT / "系统使用说明书_初版.docx"
    doc.save(path)
    return path


def build_division_table() -> Path:
    doc = Document()
    set_default_font(doc)
    add_title(doc, "小组分工确认表")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["成员", "主要职责", "具体工作内容", "备注"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    rows = [
        ["成员A", "数据处理", "整理 Excel 数据、完成数据清洗、生成赔付进展三角", "可替换为真实姓名"],
        ["成员B", "模型实现", "实现 Chain Ladder、ELR、Bornhuetter-Ferguson 模型并复核公式", "可替换为真实姓名"],
        ["成员C", "系统开发", "搭建 Streamlit 界面、实现表格展示和图表可视化", "可替换为真实姓名"],
        ["成员D", "Agent 与报告", "设计规则型解释模块、撰写报告和使用说明书", "可替换为真实姓名"],
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    path = OUTPUT / "小组分工确认表_初版.docx"
    doc.save(path)
    return path


def main() -> None:
    paths = [build_report(), build_manual(), build_division_table()]
    for path in paths:
        print(path)


if __name__ == "__main__":
    main()
