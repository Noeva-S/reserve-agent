from __future__ import annotations

from dataclasses import asdict, is_dataclass
from datetime import datetime
from io import BytesIO
from typing import Any, Mapping

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.shared import Inches

from reserve_agent.data.loader import DataQualityReport
from reserve_agent.models.reserving import ReservingOutputs


def _format_amount(value: object) -> str:
    try:
        if pd.isna(value):
            return ""
        number = float(value)
    except Exception:
        return str(value)
    if abs(number) >= 1_000_000:
        return f"{number / 1_000_000:,.2f}m"
    if abs(number) >= 1_000:
        return f"{number / 1_000:,.1f}k"
    return f"{number:,.0f}"


def _range_text(values: list[int] | tuple[int, ...] | None) -> str:
    cleaned = [value for value in (values or []) if pd.notna(value)]
    if not cleaned:
        return "未知"
    return f"{min(cleaned)}-{max(cleaned)}"


def _as_frame(value: Any) -> pd.DataFrame:
    if value is None:
        return pd.DataFrame()
    if isinstance(value, pd.DataFrame):
        return value.copy()
    if isinstance(value, pd.Series):
        return value.reset_index()
    if is_dataclass(value):
        return pd.DataFrame([asdict(value)])
    if isinstance(value, Mapping):
        return pd.DataFrame([dict(value)])
    if isinstance(value, list):
        rows = []
        for item in value:
            if is_dataclass(item):
                rows.append(asdict(item))
            elif isinstance(item, Mapping):
                rows.append(dict(item))
            else:
                rows.append({"value": item})
        return pd.DataFrame(rows)
    return pd.DataFrame([{"value": value}])


def _add_dataframe_table(doc: Document, df: pd.DataFrame, amount_cols: set[str] | None = None, max_rows: int = 20) -> None:
    if df is None or df.empty:
        doc.add_paragraph("暂无可展示数据。")
        return
    amount_cols = amount_cols or set()
    frame = df.head(max_rows).copy()
    table = doc.add_table(rows=1, cols=len(frame.columns))
    table.style = "Table Grid"
    for cell, column in zip(table.rows[0].cells, frame.columns):
        cell.text = str(column)
    for _, row in frame.iterrows():
        cells = table.add_row().cells
        for cell, column in zip(cells, frame.columns):
            value = row[column]
            cell.text = _format_amount(value) if str(column) in amount_cols else ("" if pd.isna(value) else str(value))
    if len(df) > max_rows:
        doc.add_paragraph(f"注：表格仅展示前 {max_rows} 行，完整结果请下载 Excel。")


def _add_key_value_table(doc: Document, rows: list[tuple[str, object]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "项目"
    table.rows[0].cells[1].text = "数值"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = str(label)
        cells[1].text = str(value)


def _figure_to_png_bytes(fig: Any) -> bytes:
    buffer = BytesIO()
    fig.savefig(buffer, format="png", bbox_inches="tight")
    plt.close(fig)
    return buffer.getvalue()


def _reserve_chart(outputs: ReservingOutputs) -> bytes:
    frame = outputs.comparison
    reserve_columns = [
        column
        for column in ["Chain Ladder Reserve", "ELR Reserve", "BF Reserve", "Mack Reserve"]
        if column in frame.columns
    ]
    if not reserve_columns:
        raise ValueError("No reserve columns are available.")
    plot_frame = frame[["Accident Year", *reserve_columns]].copy()
    plot_frame = plot_frame.melt("Accident Year", var_name="Method", value_name="Reserve")
    years = frame["Accident Year"].astype(int).astype(str).tolist()
    fig, ax = plt.subplots(figsize=(8, 3.8), dpi=160)
    methods = reserve_columns
    x = list(range(len(years)))
    width = 0.8 / max(len(methods), 1)
    for offset, method in enumerate(methods):
        values = plot_frame.loc[plot_frame["Method"] == method, "Reserve"].to_numpy()
        positions = [item + (offset - (len(methods) - 1) / 2) * width for item in x]
        ax.bar(positions, values, width=width, label=method.replace(" Reserve", ""))
    ax.set_xticks(x, years, rotation=45, ha="right")
    ax.set_ylabel("Reserve")
    ax.set_title("Reserve by model and accident year")
    ax.grid(axis="y", alpha=0.2)
    ax.legend(frameon=False, fontsize=7)
    fig.tight_layout()
    return _figure_to_png_bytes(fig)


def _ultimate_chart(outputs: ReservingOutputs) -> bytes:
    frame = outputs.comparison
    years = frame["Accident Year"].astype(int).tolist()
    ultimate_columns = [
        column
        for column in ["Chain Ladder Ultimate", "Expected Ultimate Loss", "BF Ultimate Loss", "Mack Ultimate"]
        if column in frame.columns
    ]
    if not ultimate_columns:
        raise ValueError("No ultimate columns are available.")
    fig, ax = plt.subplots(figsize=(8, 3.8), dpi=160)
    ax.plot(years, frame["Latest Cumulative"], marker="o", label="Latest cumulative")
    for column in ultimate_columns:
        ax.plot(years, frame[column], marker="o", label=column.replace(" Ultimate Loss", "").replace(" Ultimate", ""))
    ax.set_xlabel("Accident year")
    ax.set_ylabel("Amount")
    ax.set_title("Latest cumulative vs model ultimates")
    ax.grid(alpha=0.2)
    ax.legend(frameon=False)
    fig.tight_layout()
    return _figure_to_png_bytes(fig)


def _factor_chart(outputs: ReservingOutputs) -> bytes:
    factors = outputs.selected_factors
    fig, ax = plt.subplots(figsize=(8, 3.6), dpi=160)
    ax.plot(list(factors.index), factors.to_numpy(), marker="o")
    ax.set_xlabel("Development age")
    ax.set_ylabel("Selected factor")
    ax.set_title("Selected development factors")
    ax.grid(alpha=0.2)
    fig.tight_layout()
    return _figure_to_png_bytes(fig)


def _mack_interval_chart(outputs: ReservingOutputs) -> bytes:
    frame = outputs.mack
    if frame is None or frame.empty:
        raise ValueError("Mack results are unavailable.")
    years = frame["Accident Year"].astype(int).tolist()
    fig, ax = plt.subplots(figsize=(8, 3.8), dpi=160)
    ax.plot(years, frame["Mack Reserve"], marker="o", label="Mack reserve")
    if {"Mack 95% Lower", "Mack 95% Upper"}.issubset(frame.columns):
        ax.fill_between(
            years,
            pd.to_numeric(frame["Mack 95% Lower"], errors="coerce"),
            pd.to_numeric(frame["Mack 95% Upper"], errors="coerce"),
            alpha=0.18,
            label="95% interval",
        )
    ax.set_xlabel("Accident year")
    ax.set_ylabel("Reserve")
    ax.set_title("Mack reserve interval")
    ax.grid(alpha=0.2)
    ax.legend(frameon=False)
    fig.tight_layout()
    return _figure_to_png_bytes(fig)


def _expected_lr_sensitivity_chart(outputs: ReservingOutputs) -> bytes:
    frame = outputs.expected_lr_sensitivity
    if frame is None or frame.empty:
        raise ValueError("Expected LR sensitivity is unavailable.")
    fig, ax = plt.subplots(figsize=(8, 3.8), dpi=160)
    x = frame["Expected Loss Ratio"]
    for column in ["ELR Reserve", "BF Reserve"]:
        if column in frame.columns:
            ax.plot(x, frame[column], marker="o", label=column)
    ax.set_xlabel("Expected loss ratio")
    ax.set_ylabel("Reserve")
    ax.set_title("Expected loss ratio sensitivity")
    ax.grid(alpha=0.2)
    ax.legend(frameon=False)
    fig.tight_layout()
    return _figure_to_png_bytes(fig)


def _factor_sensitivity_chart(outputs: ReservingOutputs) -> bytes:
    frame = outputs.factor_sensitivity
    if frame is None or frame.empty:
        raise ValueError("Development factor sensitivity is unavailable.")
    fig, ax = plt.subplots(figsize=(8, 3.8), dpi=160)
    ax.plot(frame["Factor Shock"], frame["Chain Ladder Reserve"], marker="o")
    ax.set_xlabel("Factor shock")
    ax.set_ylabel("Chain Ladder reserve")
    ax.set_title("Development factor sensitivity")
    ax.grid(alpha=0.2)
    fig.tight_layout()
    return _figure_to_png_bytes(fig)


def _triangle_heatmap(triangle: pd.DataFrame) -> bytes:
    fig, ax = plt.subplots(figsize=(8, 4.2), dpi=160)
    numeric = triangle.apply(pd.to_numeric, errors="coerce")
    image = ax.imshow(numeric.to_numpy(dtype=float), aspect="auto")
    ax.set_xticks(range(len(numeric.columns)), [str(col) for col in numeric.columns])
    ax.set_yticks(range(len(numeric.index)), [str(idx) for idx in numeric.index])
    ax.set_xlabel("Development age")
    ax.set_ylabel("Accident year")
    ax.set_title("Cumulative loss triangle heatmap")
    fig.colorbar(image, ax=ax, fraction=0.046, pad=0.04)
    fig.tight_layout()
    return _figure_to_png_bytes(fig)


def build_report_chart_images(triangle: pd.DataFrame | None, outputs: ReservingOutputs | None) -> dict[str, bytes]:
    """Create PNG chart assets used by the Word report and ZIP package."""

    images: dict[str, bytes] = {}
    if triangle is not None and not triangle.empty:
        try:
            images["charts/triangle_heatmap.png"] = _triangle_heatmap(triangle)
        except Exception:
            pass
    if outputs is not None:
        for name, builder in {
            "charts/reserve_by_accident_year.png": _reserve_chart,
            "charts/latest_vs_ultimate.png": _ultimate_chart,
            "charts/development_factors.png": _factor_chart,
            "charts/mack_interval.png": _mack_interval_chart,
            "charts/expected_lr_sensitivity.png": _expected_lr_sensitivity_chart,
            "charts/factor_sensitivity.png": _factor_sensitivity_chart,
        }.items():
            try:
                images[name] = builder(outputs)
            except Exception:
                pass
    return images


def _detection_rows(
    *,
    source_file: str,
    sheet_name: str,
    format_name: str,
    measure: str,
    recognition_source: str = "",
    recognition_reason: str = "",
    header_row: int | None = None,
) -> list[tuple[str, object]]:
    rows: list[tuple[str, object]] = []
    if source_file:
        rows.append(("数据文件", source_file))
    if sheet_name:
        rows.append(("工作表", sheet_name))
    if format_name:
        rows.append(("识别格式", format_name))
    if header_row is not None:
        rows.append(("表头行", f"Excel 第 {header_row + 1} 行（内部索引 {header_row}）"))
    if recognition_source:
        rows.append(("识别来源", "DeepSeek API" if recognition_source == "api" else "本地规则"))
    if recognition_reason:
        rows.append(("识别说明", recognition_reason))
    rows.append(("建模口径", measure))
    return rows


def build_word_report(
    explanation_text: str | None = None,
    *,
    source_file: str = "",
    sheet_name: str = "",
    format_name: str = "",
    measure: str = "Paid",
    triangle: pd.DataFrame | None = None,
    outputs: ReservingOutputs | None = None,
    quality: DataQualityReport | None = None,
    method_notes: Mapping[str, str] | None = None,
    detection_summary: Any | None = None,
    validation_issues: Any | None = None,
    recognition_source: str = "",
    recognition_reason: str = "",
    header_row: int | None = None,
) -> bytes:
    """Build a complete editable Word report for the current result."""

    document = Document()
    document.add_heading("准备金评估完整报告", level=0)
    document.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")

    document.add_heading("1. 数据识别结果", level=1)
    _add_key_value_table(
        document,
        _detection_rows(
            source_file=source_file,
            sheet_name=sheet_name,
            format_name=format_name,
            measure=measure,
            recognition_source=recognition_source,
            recognition_reason=recognition_reason,
            header_row=header_row,
        ),
    )

    if isinstance(detection_summary, Mapping) and detection_summary.get("candidates"):
        candidate_rows = []
        for item in detection_summary.get("candidates", [])[:5]:
            if isinstance(item, Mapping):
                candidate_rows.append(
                    {
                        "Index": item.get("candidate_index"),
                        "Header Row": item.get("header_row_excel"),
                        "Rule Format": item.get("rule_format"),
                        "Rule Score": item.get("rule_score"),
                        "Rows": item.get("row_count"),
                        "Columns": ", ".join(map(str, item.get("candidate_columns", [])[:8])),
                    }
                )
        if candidate_rows:
            document.add_paragraph("候选表结构摘要：")
            _add_dataframe_table(document, pd.DataFrame(candidate_rows), max_rows=5)

    if quality and outputs is not None:
        document.add_heading("2. 评估概览", level=1)
        diag = outputs.diagnostics
        ay_range = _range_text(quality.accident_years)
        document.add_paragraph(
            f"系统读取 {quality.row_count:,} 行数据，事故年范围为 {ay_range}。"
            f"累计已观察赔款约 {_format_amount(diag['total_latest'])}，"
            f"Chain Ladder 准备金约 {_format_amount(diag['total_cl_reserve'])}，"
            f"BF 准备金约 {_format_amount(diag['total_bf_reserve'])}，"
            f"Mack 准备金约 {_format_amount(diag.get('total_mack_reserve', 0.0))}。"
        )
        _add_key_value_table(
            document,
            [
                ("事故年数量", f"{len(quality.accident_years):,}"),
                ("累计已观察赔款", _format_amount(diag["total_latest"])),
                ("Chain Ladder 准备金", _format_amount(diag["total_cl_reserve"])),
                ("Expected Loss Ratio 准备金", _format_amount(diag["total_elr_reserve"])),
                ("Bornhuetter-Ferguson 准备金", _format_amount(diag["total_bf_reserve"])),
                ("Mack 准备金", _format_amount(diag.get("total_mack_reserve", 0.0))),
            ],
        )

        document.add_heading("3. 数据质量诊断", level=1)
        valuation_range = _range_text(quality.valuation_years)
        _add_key_value_table(
            document,
            [
                ("源数据行数", f"{quality.row_count:,}"),
                ("赔案/事故年数量", f"{quality.claim_count:,}"),
                ("事故年范围", ay_range),
                ("评估年范围", valuation_range),
                ("缺失值数量", f"{quality.missing_values:,}"),
                ("负金额单元格", f"{quality.negative_amount_cells:,}"),
                ("零赔案/零事故年记录", f"{quality.zero_claim_rows:,}"),
            ],
        )
        for note in quality.notes:
            document.add_paragraph(note)

    issues_frame = _as_frame(validation_issues)
    if not issues_frame.empty:
        document.add_heading("4. 结构校验", level=1)
        _add_dataframe_table(document, issues_frame, max_rows=20)

    if triangle is not None:
        document.add_heading("5. 累计赔款三角形", level=1)
        triangle_frame = triangle.copy().reset_index()
        triangle_frame.columns = ["Accident Year"] + [f"Dev {int(col)}" for col in triangle.columns]
        _add_dataframe_table(document, triangle_frame, set(triangle_frame.columns[1:]))

    if outputs is not None:
        document.add_heading("6. 模型结果", level=1)
        factor_frame = outputs.selected_factors.reset_index()
        factor_frame.columns = ["Development Age", "Selected Factor"]
        _add_dataframe_table(document, factor_frame, max_rows=30)

        comparison = outputs.comparison.copy()
        _add_dataframe_table(
            document,
            comparison,
            {
                "Latest Cumulative",
                "Chain Ladder Ultimate",
                "Chain Ladder Reserve",
                "Expected Ultimate Loss",
                "ELR Reserve",
                "BF Ultimate Loss",
                "BF Reserve",
                "Mack Ultimate",
                "Mack Reserve",
                "Mack Standard Error",
                "Mack 95% Lower",
                "Mack 95% Upper",
            },
        )

        if outputs.mack is not None and not outputs.mack.empty:
            document.add_heading("6.1 Mack Chain Ladder 不确定性分析", level=2)
            mack_diag = outputs.mack_diagnostics or {}
            _add_key_value_table(
                document,
                [
                    ("Mack 准备金", _format_amount(mack_diag.get("total_mack_reserve", 0.0))),
                    ("Mack 标准误", _format_amount(mack_diag.get("total_mack_standard_error", 0.0))),
                    ("Mack CV", f"{float(mack_diag.get('total_mack_cv', 0.0)):.1%}"),
                    (
                        "95% 准备金区间",
                        f"{_format_amount(mack_diag.get('mack_95_lower', 0.0))}-"
                        f"{_format_amount(mack_diag.get('mack_95_upper', 0.0))}",
                    ),
                ],
            )
            _add_dataframe_table(
                document,
                outputs.mack,
                {
                    "Latest Cumulative",
                    "Mack Ultimate",
                    "Mack Reserve",
                    "Mack Standard Error",
                    "Mack 75% Lower",
                    "Mack 75% Upper",
                    "Mack 95% Lower",
                    "Mack 95% Upper",
                },
            )

        if outputs.expected_lr_sensitivity is not None and not outputs.expected_lr_sensitivity.empty:
            document.add_heading("6.2 期望赔付率敏感性分析", level=2)
            _add_dataframe_table(
                document,
                outputs.expected_lr_sensitivity,
                {"ELR Reserve", "ELR Ultimate", "BF Reserve", "BF Ultimate"},
                max_rows=30,
            )

        if outputs.factor_sensitivity is not None and not outputs.factor_sensitivity.empty:
            document.add_heading("6.3 发展因子敏感性分析", level=2)
            _add_dataframe_table(
                document,
                outputs.factor_sensitivity,
                {"Chain Ladder Reserve", "Chain Ladder Ultimate"},
                max_rows=30,
            )

        document.add_heading("7. 图表", level=1)
        chart_images = build_report_chart_images(triangle, outputs)
        captions = {
            "charts/triangle_heatmap.png": "图 1  累计赔款三角热力图",
            "charts/reserve_by_accident_year.png": "图 2  各模型事故年准备金",
            "charts/latest_vs_ultimate.png": "图 3  已观察赔款与各模型最终赔款",
            "charts/development_factors.png": "图 4  发展因子走势",
            "charts/mack_interval.png": "图 5  Mack 准备金区间",
            "charts/expected_lr_sensitivity.png": "图 6  期望赔付率敏感性",
            "charts/factor_sensitivity.png": "图 7  发展因子敏感性",
        }
        for filename, image_bytes in chart_images.items():
            document.add_picture(BytesIO(image_bytes), width=Inches(6.2))
            document.add_paragraph(captions.get(filename, filename))

    if explanation_text:
        document.add_heading("8. Agent 分析", level=1)
        for raw_line in explanation_text.splitlines():
            line = raw_line.strip()
            if line:
                document.add_paragraph(line[2:] if line.startswith("- ") else line)

    if method_notes:
        document.add_heading("9. 方法说明", level=1)
        for name, note in method_notes.items():
            document.add_paragraph(f"{name}: {note}")

    document.add_heading("10. 注意事项", level=1)
    document.add_paragraph(
        "本报告由课程项目中的自动化系统生成，适用于课堂展示和方法说明，不构成正式精算签字意见。"
        "实际使用前应复核数据口径、大额赔案、尾部发展、暴露数据、通胀/贴现假设和模型选择依据。"
    )

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
